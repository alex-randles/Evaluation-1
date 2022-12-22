import os
from docx import Document
from rdflib import *
import docx


# code which creates a table detailing quality issues detected in experiment 1


MQV = Namespace("http://mappingQualityVocabulary.org/")
document = Document()
table = document.add_table(rows=0, cols=6)
row_cells = table.add_row().cells
row_cells[0].text = "#"
row_cells[1].text = "Source"
row_cells[2].text = "Total Issues"
row_cells[3].text = "Description of Issue"
row_cells[4].text = "Mapping"
row_cells[5].text = "Report"


mapping_descriptions = {
    "Student Project": ["group-03-r2rml", "group-01-r2rml"],
    "OSi": ["OSi-Mapping1", "OSi-Mapping2"],
    "GeoHive": ["Townland_20m", "LEA_20m"],
    "FAIRVASC": ["FAIRVASC-Mapping3"],
}

def add_prefix(iri):
    for ns_prefix, namespace in g.namespaces():
        if iri.startswith(namespace):
            result = ns_prefix + ":" + iri.split(namespace)[-1]
            return result
    return iri


directory = os.fsencode(".")
for file in os.listdir(directory):
    mapping_directory = os.fsdecode(file)
    mapping_id = mapping_directory.split("(")[0]
    mapping_project = mapping_directory.split("(")[-1].split(")")[0]
    if os.path.isdir(file):
        row_cells = table.add_row().cells
        row_cells[0].text = mapping_id
        project_name = [name for name, mappings in mapping_descriptions.items() if mapping_project in mappings]
        if project_name:
            description = "{} ({})".format(project_name[0], mapping_project)
        else:
            description = mapping_project
        row_cells[1].text = description
        sub_directory = os.fsencode(file)
        str_sub_directory =  sub_directory.decode('utf-8')
        for sub_directory_file in os.listdir(sub_directory):
            str_file_name = sub_directory_file.decode('utf-8')
            validation_report = "./" + str_sub_directory + "/validation_report.ttl"
            if str_file_name != "validation_report.ttl" and str_file_name != "refined_mapping.ttl" and str_file_name.endswith(".ttl"):
                mapping_file = os.path.basename(sub_directory_file).decode('utf-8')
                mapping_path = str_sub_directory + "/" + mapping_file
                print(validation_report)
                violation_count = 0
                violation_descriptions = []
                if os.path.exists(validation_report):
                    g = Graph().parse(validation_report, format="ttl")
                    for s, p, o in g.triples((None, RDF.type, MQV.MappingViolation)):
                        metric = list(g.objects(s, MQV.isDescribedBy))[0]
                        metric_id = metric.split("c")[-1]
                        violation_value = list(g.objects(s, MQV.hasValue))[0]
                        prefixed_value = add_prefix(violation_value)
                        violation_count += 1
                        violation_descriptions.append("{} ({})".format(prefixed_value, metric_id))
                else:
                    print("Does not exist")
                if violation_descriptions:
                    row_cells[3].text = "•" + "\n\n •".join(violation_descriptions)
                else:
                    row_cells[3].text = "N/a"
                row_cells[2].text = str(violation_count)
                p = document.add_paragraph()
                mapping_url = "https://github.com/alex-randles/Evaluation-1/blob/main/" + mapping_path
                validation_report_url = "https://github.com/alex-randles/Evaluation-1/blob/main/" + str_sub_directory + "/validation_report.ttl"
                row_cells[4].text = mapping_url
                row_cells[5].text = validation_report_url

document.save("experiment_1_results.docx")
